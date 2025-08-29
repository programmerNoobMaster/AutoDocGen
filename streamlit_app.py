import os
import re
import io
import sys
import subprocess
from pathlib import Path
import streamlit as st

# =============== Page config ===============
st.set_page_config(page_title="Auto Doc Gen — Streamlit", layout="wide")

# =============== Session state ===============
if "stage" not in st.session_state:
    st.session_state.stage = "input"   # "input" | "running" | "results"
if "repo_url" not in st.session_state:
    st.session_state.repo_url = ""
if "wrote_files" not in st.session_state:
    st.session_state.wrote_files = []
if "combined_md" not in st.session_state:
    st.session_state.combined_md = ""

# =============== Markdown cleaning helpers ===============
HEADING_RE = re.compile(r'^\s*(#{1,6})\s+(.*\S)\s*$', re.M)
LEADING_H_RE = re.compile(r'^\s*(#{1,6})\s+(.*\S)\s*$')

def _collapse_leading_duplicate_headings(md: str) -> str:
    lines = md.splitlines()
    i = 0
    while i < len(lines) and lines[i].strip() == "":
        i += 1
    if i >= len(lines): return md
    m = LEADING_H_RE.match(lines[i])
    if not m: return md
    first_h = m.group(2).strip().lower()
    j, drop = i + 1, []
    while j < len(lines):
        if lines[j].strip() == "": drop.append(j); j += 1; continue
        mj = LEADING_H_RE.match(lines[j])
        if mj and mj.group(2).strip().lower() == first_h: drop.append(j); j += 1; continue
        break
    if not drop: return md
    keep = [line for k, line in enumerate(lines) if k not in set(drop)]
    return "\n".join(keep)

def _dedupe_consecutive_headings(md: str) -> str:
    lines, out, last = md.splitlines(), [], None
    for line in lines:
        m = HEADING_RE.match(line)
        if m:
            text = m.group(2).strip().lower()
            if last and text == last: continue
            last = text
        else:
            last = None
        out.append(line)
    return "\n".join(out)

def _strip_slug_prefix(md: str, stem: str) -> str:
    lines = md.splitlines()
    i = 0
    while i < len(lines) and lines[i].strip() == "":
        i += 1
    if i < len(lines):
        first = lines[i].strip()
        if re.fullmatch(r"[a-z0-9_]+", first) or first.replace(" ", "_").lower() == stem.lower():
            del lines[i]
    return "\n".join(lines)

def _squeeze_blank_lines(md: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", md).strip()

def clean_section_md(md: str, stem: str) -> str:
    md = _collapse_leading_duplicate_headings(md)
    md = _strip_slug_prefix(md, stem)
    md = _dedupe_consecutive_headings(md)
    md = _squeeze_blank_lines(md)
    return md

def starts_with_heading(md: str) -> bool:
    for line in md.splitlines():
        if line.strip() == "": continue
        return HEADING_RE.match(line) is not None
    return False

def humanize_stem(stem: str) -> str:
    words = re.split(r"[_\-]+", stem)
    return " ".join(w.capitalize() for w in words if w)

def combine_markdown(files):
    parts = []
    for p in files:
        p = Path(p)
        raw = p.read_text(encoding="utf-8", errors="ignore")
        cleaned = clean_section_md(raw, p.stem)
        if starts_with_heading(cleaned):
            parts.append(cleaned)
        else:
            parts.append(f"# {humanize_stem(p.stem)}\n\n{cleaned}")
    return "\n\n---\n\n".join(parts)

def naive_markdown_to_docx(md_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    for line in md_text.splitlines():
        if line.startswith("# "):
            r = doc.add_paragraph().add_run(line[2:].strip()); r.bold = True; r.font.size = Pt(16)
        elif line.startswith("## "):
            r = doc.add_paragraph().add_run(line[3:].strip()); r.bold = True; r.font.size = Pt(14)
        elif line.startswith("### "):
            r = doc.add_paragraph().add_run(line[4:].strip()); r.bold = True; r.font.size = Pt(12)
        elif line.strip().startswith(("-", "*")):
            doc.add_paragraph(line.lstrip("-* ").strip(), style="List Bullet")
        elif line.strip().startswith("```"):
            doc.add_paragraph("")  # ignore fence markers
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =============== Subprocess launcher ===============
def run_main_and_collect(repo_url: str):
    """
    Calls app/main.py --repo <url>, streams logs, collects 'Wrote: <path>' lines.
    Uses project venv Python if available. Forces UTF-8 to avoid Windows emoji issues.
    """
    base_dir = Path(__file__).parent.resolve()
    app_dir = (base_dir / "app").resolve()
    main_py = (app_dir / "main.py").resolve()
    if not main_py.exists():
        raise FileNotFoundError(f"Cannot find main.py at: {main_py}")

    venv_py = base_dir / ".venv" / "Scripts" / "python.exe"  # Windows
    py_exec = str(venv_py) if venv_py.exists() else sys.executable

    cmd = [py_exec, "-u", str(main_py), "--repo", repo_url]

    env = os.environ.copy()
    env.setdefault("PYTHONIOENCODING", "utf-8")
    env.setdefault("PYTHONUTF8", "1")

    proc = subprocess.Popen(
        cmd,
        cwd=str(base_dir),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        bufsize=1,
        env=env,
    )

    wrote_paths, log_lines = [], []
    wrote_re = re.compile(r"^Wrote:\s*(.+)$")

    # Stream logs to the page
    log_placeholder = st.empty()
    for line in proc.stdout:
        log_lines.append(line.rstrip("\n"))
        log_placeholder.code("\n".join(log_lines[-400:]))
        m = wrote_re.match(line.strip())
        if m:
            p = Path(m.group(1)).expanduser()
            if not p.is_absolute():
                p = (base_dir / p).resolve()
            wrote_paths.append(p)

    ret = proc.wait()
    if ret != 0:
        raise RuntimeError(f"main.py exited with code {ret}")

    return wrote_paths

# =============== Views ===============
def show_input_view():
    st.markdown(
        """
        <style>
        .center-box {display:flex; justify-content:center; align-items:center; margin-top:10vh;}
        .box-inner {width:min(900px, 92vw);}
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="center-box"><div class="box-inner">', unsafe_allow_html=True)

    # A form so Enter key submits; button styled as an arrow.
    with st.form("repo_form", clear_on_submit=False):
        url = st.text_input(
            "",
            value=st.session_state.repo_url,
            placeholder="Paste a GitHub repository URL…",
        )
        cols = st.columns([0.82, 0.18])
        with cols[0]:
            pass  # input is full-width above
        with cols[1]:
            submitted = st.form_submit_button("➡️  Enter", use_container_width=True)

    st.markdown('</div></div>', unsafe_allow_html=True)

    if submitted:
        if not url.strip():
            st.error("Please enter a GitHub repository URL.")
            return
        st.session_state.repo_url = url.strip()
        st.session_state.stage = "running"
        st.rerun()

def show_running_view():
    st.info("Working… this can take a few minutes on the first run.")
    try:
        files = run_main_and_collect(st.session_state.repo_url)
        st.session_state.wrote_files = files
        st.session_state.combined_md = combine_markdown(files)
        st.session_state.stage = "results"
        st.rerun()
    except Exception as e:
        import traceback
        st.error("An error occurred during generation.")
        st.exception(e)
        st.code(traceback.format_exc())
        # Offer a way back
        if st.button("↩︎ Back"):
            st.session_state.stage = "input"
            st.rerun()

def show_results_view():
    st.success("✅ Documentation generated.")

    # Preview
    st.subheader("Preview")
    for p in st.session_state.wrote_files:
        p = Path(p)
        raw = p.read_text(encoding="utf-8", errors="ignore")
        cleaned = clean_section_md(raw, p.stem)
        if not starts_with_heading(cleaned):
            st.markdown(f"# {humanize_stem(p.stem)}")
        st.markdown(cleaned)

    st.divider()

    # Downloads
    col1, col2, col3 = st.columns([0.33, 0.33, 0.34])
    with col1:
        st.download_button(
            "⬇️ Download Markdown",
            data=st.session_state.combined_md.encode("utf-8"),
            file_name="Technical_Documentation.md",
            mime="text/markdown",
            use_container_width=True,
        )
    with col2:
        try:
            docx_bytes = naive_markdown_to_docx(st.session_state.combined_md)
            st.download_button(
                "⬇️ Download DOCX (basic)",
                data=docx_bytes,
                file_name="Technical_Documentation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.info(f"DOCX export skipped: {e}")
    with col3:
        # Explicit back button so user can run another URL
        if st.button("↩︎ Back", use_container_width=True):
            st.session_state.stage = "input"
            st.session_state.repo_url = ""
            st.session_state.wrote_files = []
            st.session_state.combined_md = ""
            st.rerun()

# =============== Router ===============
st.title("🧠 Auto Doc Gen — Streamlit")  # keep title persistent

if st.session_state.stage == "input":
    show_input_view()
elif st.session_state.stage == "running":
    show_running_view()
else:
    show_results_view()
